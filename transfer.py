import psycopg2
import psycopg2.extras
import json
import datetime
import string
import os
import openpyxl
from openpyxl import Workbook
import aspose.words as aw
from flask import Flask, request, render_template,redirect, url_for, jsonify,send_file
import flask_cors
from flask_cors import CORS
import docx2txt
import docx
import re
from werkzeug.utils import secure_filename
from werkzeug.datastructures import MultiDict
app = Flask(__name__)
CORS(app)
current_time = datetime.datetime.now()
app.config['UPLOAD_FOLDER']=os.getcwd()
path = 'C:/Users/LENOVO/docgen'
file_path=''
doc_variables={'variable list':''}
file_path=''
var_list=[]
question_list=[]
temp_list=[]
var_id_list=[]
absolute_path = os.getcwd()




def get_db_connection():
    conn = psycopg2.connect(host='localhost',
                            database='docgen',
                            user='postgres',
                            password= "add@postgres")
    conn.autocommit=True
    return conn


@app.route('/upload',methods=['POST'])
def upload():
    conn = get_db_connection()
    cur =conn.cursor()
    files= request.files.getlist('attachment')
    file= files[0]
    wb = openpyxl.load_workbook(file.filename)
    sheet = wb.active
    excel_text = sheet['A1'].value
    print(excel_text)
    # text= docx2txt.process(file.filename)
    # content=docx.Document(file.filename)
    x= re.findall(r'(?<={{).*?(?=}})',excel_text)
    var_list=x
    print(var_list)
    for var in var_list:
            row = str(var)
            cur.execute('insert into variables(variable) values(%s);',(str(row),))
    file.filename = file.filename.strip('.xlsx')
    file.filename = file.filename + " " + str(current_time)+ '.docx'
    file_path = absolute_path +'/'+ str(file.filename) 
    cur.execute('insert into document(doc_name, path, timestamp) values(%s,%s,%s);',(str(file.filename),str(file_path),str(current_time)))
    filename = secure_filename(file.filename)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'],filename))
    return ['true']


@app.route('/table',methods=['POST'])
def insert_form():
    conn = get_db_connection()
    cur =conn.cursor()
    data = request.json
    q_list= data['form_q']
    if request.method == 'POST':
        cur.execute('insert into forms(form_name) values(%s);',(str(request.json['form_name']),))
        for question in q_list:
            cur.execute('insert into questions(question) values(%s);',(question,))
    return {"status":"form and question data"}

@app.route('/variable',methods=['GET'])
def index():
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    var=[]
    dict_variable = []
    cur.execute("select * from variables;")
    rows= cur.fetchall()
    for i in rows:
        dict_variable.append(dict(i))
    print(dict_variable)
    return dict_variable

@app.route('/question',methods=['GET'])
def question():
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    var=[]
    dict_question = []
    cur.execute("select * from questions;")
    rows= cur.fetchall()
    for i in rows:
        dict_question.append(dict(i))
    print(dict_question)

    return dict_question


@app.route('/insert',methods=['POST'])
def insert():
     conn = get_db_connection()
     cur = conn.cursor()
     print(request.json)
     data = request.json
     var_q_map= data['question_list']
     print(var_q_map)
     for row in var_q_map:
        var_id=str(row['var_id'])
        q_id=str(row['q_id'])
        print('the q is ',q_id)
        cur.execute('insert into mapped(var_id,q_id) values (%s,%s);',(var_id,q_id))
     return {'value':''}

    


@app.route('/retrieve',methods=['GET'])
def retrieve():
     conn = get_db_connection()
     cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
     dict_map = [] 
     question=[]
     count=0
     cur.execute("select * from mapped;")
     rows= cur.fetchall()
     for i in rows:
        dict_map.append(dict(i)) 
     for q in rows:
        cur.execute('select question from questions where q_id=%s;',(str(q[2]),))
        temp_list=cur.fetchall()
        question.append(temp_list[0][0])
     for element in dict_map:
        element['question'] = question[count]
        count=count+1
     return dict_map

@app.route('/receive',methods=['POST'])
def receive():
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    variable_list=[]
    data = request.json
    form_values= data['value_list']
    print(form_values)
    for i in form_values:
        cur.execute('select var_id from mapped where map_id=%s;',(str(i['map_id']),))
        temp_list =cur.fetchall()
        var_id_list.append(temp_list[0][0])
    print(var_id_list)
    for i in var_id_list:
         cur.execute('select variable from variables where var_id=%s;',(str(i),))
         temp_list=cur.fetchall()
         variable_list.append(temp_list[0][0])
    print(variable_list)
    filepath=os.getcwd()+'/'+'sample.docx'
    my_text = docx2txt.process(filepath)

    # print(type(my_text))
    # rep = re.sub("\{{.*?\}}", "", my_text)
    # print(rep)
    # i=0
    final = my_text
    for value in form_values:
        final= re.sub("\{{.*?\}}", value['values'], final,1)
    doc = aw.Document()
    builder = aw.DocumentBuilder(doc)
    builder.write(final)
    doc.save("out.docx")

    mydoc = docx.Document()
    mydoc.add_paragraph(final)
    path=os.getcwd()+'/'+'final_form.docx'#path with filename of generated document
    mydoc.save(path)

    return ['']
        
    
@app.route('/doc')
def doc_download():
    return render_template('download.html')


@app.route('/doc_download')
def download_file():
   path="C:/Users/LENOVO/docgen/final_form.docx"
    
   return send_file(path,as_attachment=True)






    

if __name__=='__main__':
   app.run(debug=True)

    


